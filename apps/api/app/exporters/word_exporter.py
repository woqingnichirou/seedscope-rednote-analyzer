from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from ..models import Note, Project
from .common import avg, ensure_parent, split_notes, top_counter


def _set_font(run, bold: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)


def _paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_font(run, bold=bold)


def _heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading("", level=level)
    run = heading.add_run(text)
    _set_font(run, bold=True, size=16 if level == 1 else 13)


def _table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(header)
        _set_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(value))
            _set_font(run)


def export_report_word(project: Project, notes: list[Note], path: Path) -> None:
    ensure_parent(path)
    brand_a_notes, brand_b_notes = split_notes(notes)
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    _heading(document, f"SeedScope 竞品分析报告：{project.brand_a} vs {project.brand_b}", 0)
    _paragraph(document, f"行业：{project.industry}｜周期：{project.period}")
    _paragraph(document, f"分析目标：{project.objective}")

    _heading(document, "1）核心结论")
    _paragraph(document, "结论先行：当前优化重点不是增加内容数量，而是提升决策型内容比例、证据密度和 CTA 自然承接。", bold=True)
    _paragraph(document, "Brand A 应优先补齐标题公式、封面证据位、正文判断标准和可复投素材池。")

    _heading(document, "2）样本概览")
    _table(
        document,
        ["指标", project.brand_a, project.brand_b],
        [
            ["样本数", len(brand_a_notes), len(brand_b_notes)],
            ["平均点赞", f"{avg(brand_a_notes, 'likes'):.1f}", f"{avg(brand_b_notes, 'likes'):.1f}"],
            ["平均收藏", f"{avg(brand_a_notes, 'collects'):.1f}", f"{avg(brand_b_notes, 'collects'):.1f}"],
            ["平均评论", f"{avg(brand_a_notes, 'comments'):.1f}", f"{avg(brand_b_notes, 'comments'):.1f}"],
        ],
    )

    _heading(document, "3）发布时间节奏")
    _paragraph(document, "建议按关键决策窗口集中发布同一主题的多角度内容，避免主题分散和平均铺量。", bold=True)
    _paragraph(document, "后续可结合发布时间字段，按月度/周度识别内容波峰和投放节奏。")

    _heading(document, "4）标题策略对比")
    _table(document, ["标题类型", "数量"], [[k, v] for k, v in top_counter(notes, "title_type")])
    _paragraph(document, "优先增加问题先行、方法攻略和对比测评标题，减少品牌前置和口号型表达。", bold=True)

    _heading(document, "5）封面策略对比")
    _table(document, ["封面类型", "数量"], [[k, v] for k, v in top_counter(notes, "cover_type")])
    _paragraph(document, "封面应同时具备阶段词、痛点词、结果词和证据位。")

    _heading(document, "6）内容母题分布")
    _table(document, ["内容类型", "数量"], [[k, v] for k, v in top_counter(notes, "content_type")])
    _paragraph(document, "高赞内容应沉淀为可复用母题，而不是停留在单篇复盘。")

    _heading(document, "7）卖点表达对比")
    _paragraph(document, "Brand A 需要把产品卖点翻译成用户选择理由，避免只堆叠品牌优势。", bold=True)
    _paragraph(document, "建议表达路径：具体痛点 → 判断标准 → 品牌如何满足标准 → 结果证据。")

    _heading(document, "8）CTA承接分析")
    _table(document, ["CTA 类型", "数量"], [[k, v] for k, v in top_counter(notes, "cta_type")])
    _paragraph(document, "CTA 应放在正文后半段，以评论、私信、资料包或试听观察自然承接。")

    _heading(document, "9）本品问题总结")
    for item in ["高决策作者占比不足", "品牌前置标题偏多", "封面证据位不足", "正文方法论段落偏弱", "高效素材复投机制不足"]:
        _paragraph(document, f"• {item}")

    _heading(document, "10）下阶段策略建议")
    for item in ["重构达人池：高决策作者、结果陪跑作者、低成本测试作者分层管理", "重建模板库：标题、封面、正文、CTA 标准化", "建立素材白名单：按收藏、评论、转化反馈筛选复投素材", "周度复盘：按达人 × 选题 × 封面 × CTA 归因"]:
        _paragraph(document, f"• {item}")

    document.save(path)
